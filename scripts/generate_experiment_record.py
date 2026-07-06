"""生成 docs/实验记录.docx — RouteAlpha 实验目的、迭代提升、结论与未来方向。

用法: python scripts/generate_experiment_record.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "实验记录.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    title = doc.add_heading("RouteAlpha 实验记录", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "一、实验目的", 1)
    add_para(
        doc,
        "在大模型生态中, 不同模型在能力、延迟与成本上差异显著。"
        "若对所有 query 一律调用最强/最贵模型, 成本不可控; 若一律用最便宜模型, 复杂任务成功率不足。"
        "本项目的核心目标是: 在全局预算硬约束下, 自动为每条 query 选择最合适的模型, "
        "在同等成本下提高成功率, 或在同等质量下显著节约成本 (predict-then-optimize)。",
    )
    add_para(doc, "技术路线与电力交易项目同构: XGB 预测每条 query 在各模型上的成功率 → Gurobi MILP 在预算内做全局最优分配 → 滚动回测诚实评估。", bold=False)

    add_heading(doc, "二、系统架构", 1)
    add_bullet(doc, "预测层: bge-small embedding + 结构特征 → 每 model 独立 XGBoost → Isotonic 概率校准")
    add_bullet(doc, "决策层: Gurobi MILP, 约束 Σ cost ≤ Budget, 目标 max Σ P(success)")
    add_bullet(doc, "评测层: OOF 回测 (accuracy/AUC/Brier/ECE)、Pareto 前沿、五类失败归因、LLM-as-judge 偏置诊断")

    add_heading(doc, "三、关键迭代与提升", 1)
    add_para(doc, "数据集: RouterBench peek 1000 条 / 900 query 测试集 (OOF)。特征与校准迭代如下:", bold=True)

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["版本", "特征", "ECE", "MILP SR@0.002", "MILP SR@0.001", "gap@0.002"]):
        hdr[i].text = h

    rows = [
        ("v0 baseline", "bge384 + 24维结构, 无TE", "0.0565", "0.802", "0.762", "0.124"),
        ("v2 TE only", "v0 + Target Encoding (α=10)", "0.0542", "0.799", "0.769", "0.128"),
        ("v3 TE+cross", "v2 + cross difficulty + mglobal (α=20)", "0.0516", "0.810", "0.777", "0.117"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val

    add_para(doc, "")
    add_para(doc, "概率校准 (v3 最终配置):", bold=True)
    add_bullet(doc, "ECE raw → cal: 0.184 → 0.052 (Isotonic, fold 内 holdout)")
    add_bullet(doc, "校准前模型过度自信, 优化器用未校准概率下注会系统性选错; 校准后可靠性图贴近对角线")

    add_heading(doc, "四、最终结论 (v3)", 1)
    add_bullet(doc, "预测 overall AUC: 0.643")
    add_bullet(doc, "MILP @0.002/query: 真实成功率 0.810, optimality gap 0.117 (oracle 0.927)")
    add_bullet(doc, "相对 always-expensive: 成本约 60%, 质量保持约 95%")
    add_bullet(doc, "MILP @0.001/query (紧预算): 成功率 0.777, 路由增益最明显")
    add_bullet(doc, "失败归因 (171 条): 校准/概率误导 82, 任务本身难 66, 预测排序错 8 — 说明下一刀应在校准与难例, 而非单纯加特征")
    add_bullet(doc, "LLM-as-judge (Kimi moonshot-v1-8k, 12 对): flip rate 0.083, accept 0.917, 采纳准确率 1.0; hard 子集暴露 1 条长度偏置翻转并被 swap-and-aggregate 丢弃")

    add_heading(doc, "五、数据纪律", 1)
    add_bullet(doc, "四分法: train / calibration / test / golden — golden 永不参与训练")
    add_bullet(doc, "修复特征穿越: TF-IDF 由全量 fit 改为每 fold 仅 train fit")
    add_bullet(doc, "诚实评估: MILP 真实成功率用 y_true, 不用预测值")

    add_heading(doc, "六、未来可提升方向", 1)
    add_bullet(doc, "级联路由 (cascade): 先 cheap 试答, 低置信再升级强模型 — 与单跳 MILP 互补")
    add_bullet(doc, "RouterArena held-out 打榜: 用真实 API 验证泛化, 拿可验证排名")
    add_bullet(doc, "更强特征: 任务类型 embedding、query 难度估计、历史成功率时序")
    add_bullet(doc, "风险约束 MILP: 除期望成功外加入方差/最坏情况惩罚 (量化交易味)")
    add_bullet(doc, "扩大样本: peek 1000 → 全量 RouterBench / 时序滚动回测")

    doc.save(OUT)
    print(f"已写入: {OUT}")


if __name__ == "__main__":
    main()
